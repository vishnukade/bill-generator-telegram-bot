import os
import subprocess
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes, ConversationHandler, CallbackQueryHandler
from docx import Document
import nest_asyncio 
# Define categories of placeholders
INPUT_GROUPS = [
    {"message": "✍️ Enter the client information (Client Name, Client Mobile no, Client Address, Company) separated by commas.",
     "keys": ["ClientName", "clientNo", "Address", "company"]},
    {"message": "🚗 Enter car details (Car Name, AC or Non-AC, Car Registration No, Driver Name) separated by commas.",
     "keys": ["CarName", "AcN", "carNo", "Dname"]},
    {"message": "📅 Enter the trip details (Start Date, Close Date, Total Days**) separated by commas.",
     "keys": ["Sdate", "CDate", "Ndays"]},
    {"message": "⏰ Enter time (Start Time, Close Time, Total Time) separated by commas.",
     "keys": ["Stime", "Ctime", "Ttime"]},
    {"message": "📍** Enter the place**",
     "keys": ["Place"]},
    {"message": "🛣️ Enter KM (Start KM , Close KM , Total KM) separated by commas.",
     "keys": ["Skm", "Ckm", "Tkm"]},
    {"message": "💰 Enter (Package, Rate per KM, Amount) separated by commas.",
     "keys": ["Package", "RatePerKm", "Amount"]},
    {"message": "🚕 Enter (Toll and Parking, Total Amount, Balance) separated by commas.",
     "keys": ["TollPark", "Total", "Balance"]},
    {"message": "💵 Enter the amount in words",
     "keys": ["AmountWord"]},
]

# States for ConversationHandler
ASKING, GENERATING = range(2)

# Dictionary to store user inputs
user_data = {}
template_choice = ""

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Start command to initiate the bot."""
    # Send welcome message with buttons
    keyboard = [
        [
            InlineKeyboardButton("Bramhadev", callback_data="template1"),
            InlineKeyboardButton("Ganesh", callback_data="template2"),
        ]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text(
        "🎉 Welcome to the Bill Generator App! \nFollow the instructions below to generate your bill 🧾\n\nPlease choose a template:",
        reply_markup=reply_markup,
    )

    return ASKING

async def button_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle button presses."""
    query = update.callback_query
    await query.answer()  # Answer the callback query
    global template_choice

    # Set the template choice based on the button press
    if query.data == "template1":
        template_choice = "template1.docx"
        await query.edit_message_text("You selected the Bramhadev template. Please enter the details below.")
    elif query.data == "template2":
        template_choice = "template2.docx"
        await query.edit_message_text("You selected the Ganesh template. Please enter the details below.")

    # Begin asking for the first input group
    context.user_data["current_group_index"] = 0
    first_group = INPUT_GROUPS[0]
    await query.message.reply_text(first_group["message"], parse_mode='Markdown')

    return ASKING

async def ask_details(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Ask user for grouped inputs."""
    group_index = context.user_data["current_group_index"]
    group = INPUT_GROUPS[group_index]
    user_input = update.message.text

    # Process user input
    values = [value.strip() for value in user_input.split(",")]
    if len(values) != len(group["keys"]):
        await update.message.reply_text(f"❌ Invalid input. Please provide {len(group['keys'])} values separated by commas.")
        return ASKING

    # Save user inputs to `user_data`
    for key, value in zip(group["keys"], values):
        user_data[key] = value

    # Special mapping:
    # Replace Date and Rtime with Stime, and TotalKm with Tkm
    if "Stime" in user_data:
        user_data["Date"] = user_data["Sdate"]  # Map Stime to Date
        user_data["Rtime"] = user_data["Stime"]  # Map Stime to Rtime
    if "Tkm" in user_data:
        user_data["TotalKm"] = user_data["Tkm"]  # Map Tkm to TotalKm

    # Move to the next group or generate the PDF
    if group_index + 1 < len(INPUT_GROUPS):
        context.user_data["current_group_index"] += 1
        next_group = INPUT_GROUPS[context.user_data["current_group_index"]]
        await update.message.reply_text(next_group["message"])
        return ASKING
    else:
        await update.message.reply_text("✅ Thank you! Generating your PDF... 📄")
        return await generate_pdf(update, context)

async def generate_pdf(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Generate PDF from template."""
    # Load the DOCX template based on the user's choice
    if template_choice == "template1.docx":
        template_path = "/content/template1.docx"  # Update with your actual template path
    elif template_choice == "template2.docx":
        template_path = "/content/template2.docx"  # Update with your actual template path
    else:
        await update.message.reply_text("❌ Template not selected properly.")
        return ConversationHandler.END

    output_docx_path = "output.docx"
    output_pdf_path = "output.pdf"

    # Replace placeholders in the template, including tables
    doc = Document(template_path)

    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for placeholder, value in user_data.items():
            if f"{{{{{placeholder}}}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{{{placeholder}}}}}", value)

    # Replace placeholders inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in user_data.items():
                    if f"{{{{{placeholder}}}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{{{placeholder}}}}}", value)

    doc.save(output_docx_path)

    try:
        # Convert DOCX to PDF using LibreOffice
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", output_docx_path],
            check=True
        )

        # Adjust the PDF file path (LibreOffice appends `.pdf`)
        if not os.path.exists(output_pdf_path):
            output_pdf_path = output_docx_path.replace(".docx", ".pdf")

        # Send the converted PDF file to the user
        with open(output_pdf_path, "rb") as pdf_file:
            await update.message.reply_document(pdf_file)

    except Exception as e:
        await update.message.reply_text(f"❌ Error generating PDF: {e}")
    finally:
        # Clean up temporary files
        if os.path.exists(output_docx_path):
            os.remove(output_docx_path)
        if os.path.exists(output_pdf_path):
            os.remove(output_pdf_path)

    return ConversationHandler.END

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Cancel the conversation."""
    await update.message.reply_text("❌ Operation canceled. Let me know if you need anything else! 😊")
    return ConversationHandler.END

def main():
    """Run the bot."""
    bot_token = "8010187997:AAF3QL16cErai9i_qzYHqCoom6UMfF-hPaM"  # Replace with your bot token
    application = Application.builder().token(bot_token).build()

    conv_handler = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={
            ASKING: [MessageHandler(filters.TEXT & ~filters.COMMAND, ask_details), CallbackQueryHandler(button_handler)],
            GENERATING: [MessageHandler(filters.TEXT & ~filters.COMMAND, generate_pdf)],
        },
        fallbacks=[CommandHandler("cancel", cancel)],
    )

    application.add_handler(conv_handler)
    nest_asyncio.apply()
    application.run_polling()

if __name__ == "__main__":
    main()
